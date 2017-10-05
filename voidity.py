#!/bin/env python3

import os
import docx
import sys
import magic
from PIL import Image
from glob import glob
from argparse import ArgumentParser
from collections import defaultdict
import json

fmtstext = ['application/msword',
            'application/pdf',
            'text/css',
            'text/html',
            'text/plain',
            'text/xml',
            'text/csv',
            'text/tab-separated-values',
            'text/rtf',
            'text/sgml'
            ]

fmtsimg =  ['image/bmp',
            'image/gif',
            'image/jp2',
            'image/jpeg',
            'image/png',
            'image/tiff',
            'image/x-jng'
            ]

fmtsau =   ['application/ogg',
            'audio/mpeg',
            'audio/midi',
            'audio/ogg',
            'audio/x-wav'
            ]

fmtsvid =  ['video/mp4',
            'video/mpeg',
            'video/ogg',
            'video/quicktime'
            ]


## GENERAL TESTS ##
# Size of file
def _size(obj):
    try:
        res = os.stat(obj)
    except:
        return None
    if res.st_size > 1000000:
        return False
    else:
        return True


## TEXT TESTS ##
def _word_length(obj):
    if len(obj.paragraphs) < 2:
        return True
    else:
        w = 0
        for o in obj.paragraphs:
            w = w + len(o.text)
        if w < 100:
            return True
        else:
            return False


def _text_length(obj):
    if len(obj) < 100:
        return True
    else:
        return False


## IMAGE TESTS ##
def _image_dimensions(obj):
    if (obj.size[0] < 100) or (obj.size[1] < 100):
        return True
    else:
        return False


def _image_color(obj):
    if len(obj.getcolors()) < 3:
        return True
    else:
        return False


## IMAGE FORMAT TESTS ##
#http://git.imagemagick.org/repos/ImageMagick/commit/501b648ee40f804228c76fddc02ca479c75666f3
def _png_min_size(obj):
    if os.path.getsize(obj) < 61:
        return True
    else:
        return False


#http://git.imagemagick.org/repos/ImageMagick/commit/f9574dc71cc1ab8219b3bdfba11bf67dc2d98c71
def _jpeg_min_size(obj):
    if os.path.getsize(obj) < 107:
        return True
    else:
        return False


#http://git.imagemagick.org/repos/ImageMagick/commit/3cc9d45352ebb92947d27c46e2604104b7ebfe90
def _jng_min_size(obj):
    if os.path.getsize(obj) < 147:
        return True
    else:
        return False

#http://git.imagemagick.org/repos/ImageMagick/commit/501b648ee40f804228c76fddc02ca479c75666f3
def png_min_size(obj):
    if os.path.getsize(obj) < 61:
        return True
    else:
        return False

#http://git.imagemagick.org/repos/ImageMagick/commit/f9574dc71cc1ab8219b3bdfba11bf67dc2d98c71
def jpeg_min_size(obj):
    if os.path.getsize(obj) < 107:
        return True
    else:
        return False

#http://git.imagemagick.org/repos/ImageMagick/commit/3cc9d45352ebb92947d27c46e2604104b7ebfe90
def jng_min_size(obj):
    if os.path.getsize(obj) < 147:
        return True
    else:
        return False

def runtests(filename):
    results = defaultdict(dict)

    fmime = magic.from_file(filename, mime=True)

    results['general']['less_than_1mb'] = _size(filename)

    if fmime in fmtsimg:
        try:
            imgfile = Image.open(filename)
        except:
            imgfile = None
        if fmime == 'image/png':
            results['images']['png_min_size'] = _png_min_size(filename)
        if fmime == 'image/jpeg':
            results['images']['jpeg_min_size'] = _jpeg_min_size(filename)
        if fmime == 'image/x-jng':
            results['images']['jng_min_size'] = _jng_min_size(filename) 
        else:
            if imgfile:
                results['images']['less_than_100x100'] = _image_dimensions(imgfile)
                results['images']['less_than_3_colors'] = _image_color(imgfile)
            else:
                results['images']['less_than_100x100'] = None
                results['images']['less_than_3_colors'] = None

    if fmime in fmtstext:
        if fmime == 'application/msword':
            results['text']['length_less_than_100c'] = _word_length(docx.Document(filename))
        else:
            results['text']['length_less_than_100c'] = _text_length(open(filename).read())
 
    return results


def main():
    parser = ArgumentParser()
    parser.add_argument('input_dir', metavar='[input directory]',
                        help='Path of files to scan')
    parser.add_argument('output', metavar='[output file]',
                        help='File to output report (will not overwrite)')
    args = parser.parse_args()

    if os.path.isfile(args.output):
        sys.exit('error: output file already exists')

    if os.path.exists(args.input_dir) and os.path.isdir(args.input_dir):
        tc = [_tc for _tc in glob(os.path.join(args.input_dir, '*')) if not os.path.isdir(_tc)]
    else:
        sys.exit('error: input directory doesn\'t exist or the input directory isn\'t a directory')

    allres = {}
    for tf in tc:

    with open(args.output, 'w') as output:
        json.dump(allres, output, indent=4, sort_keys=True) 


if __name__ == '__main__':
    main()
